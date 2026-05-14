from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

def add_divider(doc, color="1D9E75"):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(29, 158, 117)
    add_divider(doc)

def add_bullet(doc, text, font_size=10):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text.lstrip('•- ·'))
    run.font.size = Pt(font_size)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.space_before = Pt(0)

def generate_resume_docx(rewritten_data, original_data: dict, page_limit: str = "Full (no limit)") -> bytes:
    doc = Document()

    # --- Page margins based on page limit ---
    section = doc.sections[0]
    if page_limit == "1 Page":
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        font_size = 9
        bullet_limit = 2  # max 2 bullets per job
        show_projects = False
        show_achievements = False
    elif page_limit == "2 Pages":
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        font_size = 10
        bullet_limit = 4  # max 4 bullets per job
        show_projects = True
        show_achievements = True
    else:  # Full
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        font_size = 10
        bullet_limit = 999  # no limit
        show_projects = True
        show_achievements = True

    # ============================
    # HEADER — Name + Contact Info
    # ============================
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(rewritten_data.candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(18 if page_limit == "1 Page" else 22)
    name_run.font.color.rgb = RGBColor(29, 158, 117)
    name_para.paragraph_format.space_after = Pt(2)

    # Contact info
    contact_parts = []
    if original_data.get("email"):
        contact_parts.append(original_data["email"])
    if original_data.get("phone"):
        contact_parts.append(original_data["phone"])
    if original_data.get("linkedin"):
        contact_parts.append(original_data["linkedin"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        contact_run.font.size = Pt(font_size)
        contact_run.font.color.rgb = RGBColor(100, 100, 100)
        contact_para.paragraph_format.space_after = Pt(3)

    add_divider(doc)

    # ============================
    # PROFESSIONAL SUMMARY
    # ============================
    add_section_heading(doc, "Professional Summary")
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(rewritten_data.improved_summary)
    summary_run.font.size = Pt(font_size)
    summary_run.italic = True
    summary_para.paragraph_format.space_after = Pt(3)

    # ============================
    # SKILLS
    # ============================
    add_section_heading(doc, "Skills")
    skills_text = "  •  ".join(rewritten_data.improved_skills)
    skills_para = doc.add_paragraph()
    skills_run = skills_para.add_run(skills_text)
    skills_run.font.size = Pt(font_size)
    skills_para.paragraph_format.space_after = Pt(3)

    # ============================
    # WORK EXPERIENCE
    # ============================
    add_section_heading(doc, "Work Experience")
    for exp in rewritten_data.improved_experience:
        job_para = doc.add_paragraph()
        job_para.paragraph_format.space_after = Pt(1)
        job_para.paragraph_format.space_before = Pt(3)
        title_run = job_para.add_run(exp.job_title)
        title_run.bold = True
        title_run.font.size = Pt(font_size + 1)

        details_para = doc.add_paragraph()
        details_para.paragraph_format.space_after = Pt(1)
        details_para.paragraph_format.space_before = Pt(0)
        details_run = details_para.add_run(
            f"{exp.company}  |  {exp.duration}"
        )
        details_run.font.size = Pt(font_size)
        details_run.font.color.rgb = RGBColor(100, 100, 100)
        details_run.italic = True

        # Limit bullets based on page selection
        bullets = exp.improved_bullets[:bullet_limit]
        for bullet in bullets:
            add_bullet(doc, bullet, font_size)

    # ============================
    # EDUCATION
    # ============================
    education = original_data.get("education", [])
    if education:
        add_section_heading(doc, "Education")
        for edu in education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_after = Pt(1)
            edu_para.paragraph_format.space_before = Pt(3)
            degree_run = edu_para.add_run(edu.get("degree", ""))
            degree_run.bold = True
            degree_run.font.size = Pt(font_size + 1)

            school_para = doc.add_paragraph()
            school_para.paragraph_format.space_after = Pt(2)
            school_para.paragraph_format.space_before = Pt(0)
            school_run = school_para.add_run(
                f"{edu.get('institution', '')}  |  {edu.get('year', '')}"
            )
            school_run.font.size = Pt(font_size)
            school_run.font.color.rgb = RGBColor(100, 100, 100)
            school_run.italic = True

    # ============================
    # PROJECTS (controlled by page limit)
    # ============================
    projects = original_data.get("projects", [])
    if projects and show_projects:
        add_section_heading(doc, "Projects")
        for project in projects:
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(3)
            proj_para.paragraph_format.space_after = Pt(1)
            proj_run = proj_para.add_run(project.get("name", ""))
            proj_run.bold = True
            proj_run.font.size = Pt(font_size + 1)

            bullets = project.get("description", [])[:bullet_limit]
            for bullet in bullets:
                add_bullet(doc, bullet, font_size)

    # ============================
    # CERTIFICATIONS
    # ============================
    certifications = original_data.get("certifications", [])
    if certifications:
        add_section_heading(doc, "Certifications")
        for cert in certifications:
            cert_para = doc.add_paragraph()
            cert_para.paragraph_format.space_before = Pt(2)
            cert_para.paragraph_format.space_after = Pt(1)
            cert_run = cert_para.add_run(cert.get("name", ""))
            cert_run.bold = True
            cert_run.font.size = Pt(font_size)
            if cert.get("issuer") or cert.get("year"):
                issuer_run = cert_para.add_run(
                    f"  —  {cert.get('issuer', '')}  |  {cert.get('year', '')}"
                )
                issuer_run.font.size = Pt(font_size)
                issuer_run.font.color.rgb = RGBColor(100, 100, 100)

    # ============================
    # ACHIEVEMENTS (controlled by page limit)
    # ============================
    achievements = original_data.get("achievements", [])
    if achievements and show_achievements:
        add_section_heading(doc, "Achievements")
        for achievement in achievements:
            add_bullet(doc, achievement, font_size)

    # ============================
    # LANGUAGES
    # ============================
    languages = original_data.get("languages", [])
    if languages:
        add_section_heading(doc, "Languages")
        lang_para = doc.add_paragraph()
        lang_run = lang_para.add_run("  •  ".join(languages))
        lang_run.font.size = Pt(font_size)

    # --- Save to bytes ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()